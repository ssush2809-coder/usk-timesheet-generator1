from __future__ import annotations

import json
import re
import tempfile
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any, Dict, Iterable, Optional, Tuple

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches

APP_DIR = Path(__file__).resolve().parent
DEFAULT_CONFIG_PATH = APP_DIR / "default_overlay_config.json"
DATA_CONFIG_PATH = APP_DIR / "overlay_config.json"

DAYS = [
    "monday",
    "tuesday",
    "wednesday",
    "thursday",
    "friday",
    "saturday",
    "sunday",
]


def load_overlay_config() -> Dict[str, Any]:
    path = DATA_CONFIG_PATH if DATA_CONFIG_PATH.exists() else DEFAULT_CONFIG_PATH
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def save_overlay_config(config: Dict[str, Any]) -> None:
    DATA_CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(DATA_CONFIG_PATH, "w", encoding="utf-8") as f:
        json.dump(config, f, indent=2)


def reset_overlay_config() -> None:
    if DATA_CONFIG_PATH.exists():
        DATA_CONFIG_PATH.unlink()


def week_dates_from_monday(monday: date) -> Dict[str, date]:
    return {day: monday + timedelta(days=i) for i, day in enumerate(DAYS)}


def monday_for_date(d: date) -> date:
    return d - timedelta(days=d.weekday())


def format_week_range(monday: date) -> str:
    sunday = monday + timedelta(days=6)
    return f"{monday.strftime('%m/%d/%Y')} - {sunday.strftime('%m/%d/%Y')}"


def format_number(value: float) -> str:
    if float(value).is_integer():
        return str(int(value))
    return f"{value:.2f}".rstrip("0").rstrip(".")


def safe_filename(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9_\-]+", "_", value.strip())
    return value.strip("_") or "timesheet"


def _relative_rect(page: fitz.Page, box: Dict[str, float]) -> fitz.Rect:
    rect = page.rect
    x0 = rect.x0 + box["x"] * rect.width
    y0 = rect.y0 + box["y"] * rect.height
    x1 = x0 + box["w"] * rect.width
    y1 = y0 + box["h"] * rect.height
    return fitz.Rect(x0, y0, x1, y1)


def _alignment(value: str) -> int:
    value = (value or "center").lower()
    if value == "left":
        return fitz.TEXT_ALIGN_LEFT
    if value == "right":
        return fitz.TEXT_ALIGN_RIGHT
    return fitz.TEXT_ALIGN_CENTER


def _draw_text(page: fitz.Page, field: Dict[str, Any], text: str) -> None:
    if text is None:
        return
    text = str(text).strip()
    if not text:
        return
    rect = _relative_rect(page, field)
    page.insert_textbox(
        rect,
        text,
        fontsize=float(field.get("size", 10)),
        fontname=field.get("font", "tiro"),
        color=(0, 0, 0),
        align=_alignment(field.get("align", "center")),
    )


def _insert_signature(page: fitz.Page, field: Dict[str, Any], image_path: Optional[Path]) -> None:
    if not image_path:
        return
    if not image_path.exists():
        return
    rect = _relative_rect(page, field)
    page.insert_image(rect, filename=str(image_path), keep_proportion=True, overlay=True)


def payload_to_overlay_values(payload: Dict[str, Any]) -> Dict[str, str]:
    hours = payload.get("hours", {})
    dates = payload.get("dates", {})
    total_hours = sum(float(hours.get(day, 0) or 0) for day in DAYS)

    values: Dict[str, str] = {
        "employee_name": payload.get("employee_name", ""),
        "supervisor_name": payload.get("supervisor_name", ""),
        "reporting_week": payload.get("reporting_week", ""),
        "total_hours": format_number(total_hours),
    }

    for day in DAYS:
        values[f"{day}_date"] = str(dates.get(day, ""))
        values[f"{day}_hours"] = format_number(float(hours.get(day, 0) or 0))

    return values


def generate_pdf(
    *,
    template_pdf: Path,
    output_pdf: Path,
    payload: Dict[str, Any],
    signature_image: Optional[Path] = None,
    overlay_config: Optional[Dict[str, Any]] = None,
) -> Path:
    if not template_pdf.exists():
        raise FileNotFoundError(f"Template PDF not found: {template_pdf}")
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    config = overlay_config or load_overlay_config()
    fields = config.get("fields", {})
    values = payload_to_overlay_values(payload)

    doc = fitz.open(str(template_pdf))
    page_index = int(config.get("page_index", 0))
    if page_index >= len(doc):
        raise ValueError(f"Template has {len(doc)} page(s), but config page_index is {page_index}")
    page = doc[page_index]

    for name, value in values.items():
        field = fields.get(name)
        if field:
            _draw_text(page, field, value)

    signature_field = fields.get("employee_signature")
    if signature_field and signature_image:
        _insert_signature(page, signature_field, Path(signature_image))

    doc.save(str(output_pdf), garbage=4, deflate=True, clean=True)
    doc.close()
    return output_pdf


def pdf_to_word_image_docx(pdf_path: Path, docx_path: Path, dpi: int = 220) -> Path:
    """Create a Word file that visually preserves the generated PDF as full-page images.

    This is intentionally image-based because the original timesheet is a fixed-layout PDF template.
    It keeps the Word export submission-ready instead of attempting to rebuild the form in Word.
    """
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    pdf = fitz.open(str(pdf_path))
    document = Document()

    # Remove default paragraph spacing and page margins.
    for i, page in enumerate(pdf):
        if i > 0:
            document.add_page_break()

        section = document.sections[-1]
        width_inches = page.rect.width / 72
        height_inches = page.rect.height / 72
        section.page_width = Inches(width_inches)
        section.page_height = Inches(height_inches)
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        section.header_distance = Inches(0)
        section.footer_distance = Inches(0)

        zoom = dpi / 72
        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp_path = Path(tmp.name)
        pix.save(str(tmp_path))

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.line_spacing = 1
        run = paragraph.add_run()
        run.add_picture(str(tmp_path), width=Inches(width_inches))
        tmp_path.unlink(missing_ok=True)

    pdf.close()
    document.save(str(docx_path))
    return docx_path


def generate_timesheet_files(
    *,
    template_pdf: Path,
    output_dir: Path,
    payload: Dict[str, Any],
    signature_image: Optional[Path] = None,
) -> Tuple[Path, Path]:
    employee = safe_filename(payload.get("employee_name", "employee"))
    week = safe_filename(payload.get("reporting_week", "week"))
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base = f"{employee}_{week}_{timestamp}"
    pdf_path = output_dir / f"{base}.pdf"
    docx_path = output_dir / f"{base}.docx"

    generate_pdf(
        template_pdf=template_pdf,
        output_pdf=pdf_path,
        payload=payload,
        signature_image=signature_image,
    )
    pdf_to_word_image_docx(pdf_path, docx_path)
    return pdf_path, docx_path
